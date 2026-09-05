#!/usr/bin/env python3
"""把项目里的 Markdown 报告转成 Word 文档。

只支持本项目实际用到的子集：标题、段落、表格、引用块、无序/有序列表、
粗体、行内代码、分隔线、Markdown 图片。故意不做通用 Markdown 解析器。

用法：python3 scripts/md_to_docx.py 输入.md 输出.docx
"""
import os
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

CJK_FONT = "微软雅黑"


def set_cjk(run):
    run.font.name = CJK_FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), CJK_FONT)


INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")


def add_inline(paragraph, text, base_bold=False, color=None):
    """把 **粗体** 和 `行内代码` 渲染成不同的 run。"""
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.color.rgb = RGBColor(0xB0, 0x30, 0x60)
        else:
            run = paragraph.add_run(part)
            run.bold = base_bold
        if color is not None:
            run.font.color.rgb = color
        set_cjk(run)


def is_table_row(line):
    return line.startswith("|") and line.endswith("|")


def split_row(line):
    return [c.strip() for c in line.strip("|").split("|")]


def is_separator_row(line):
    return bool(re.fullmatch(r"\|[\s:\-|]+\|", line))


IMG = re.compile(r"^!\[(.*?)\]\((.+?)\)$")


def build(md_path, docx_path):
    md_dir = os.path.dirname(os.path.abspath(md_path))
    with open(md_path, encoding="utf-8") as fh:
        lines = fh.read().split("\n")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)

    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("---") and set(stripped) <= {"-"}:
            doc.add_paragraph()
            i += 1
            continue

        img = IMG.match(stripped)
        if img:
            alt, src = img.group(1), img.group(2)
            src_path = src if os.path.isabs(src) else os.path.normpath(os.path.join(md_dir, src))
            if os.path.isfile(src_path):
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(src_path, width=Inches(6.3))
                if alt:
                    cap = doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    add_inline(cap, alt, color=RGBColor(0x55, 0x55, 0x55))
            else:
                para = doc.add_paragraph()
                add_inline(para, f"[缺图 {alt or src}]", color=RGBColor(0xB0, 0x30, 0x60))
            i += 1
            continue

        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            heading = doc.add_heading(level=min(level, 4))
            add_inline(heading, m.group(2))
            for run in heading.runs:
                set_cjk(run)
            i += 1
            continue

        # 表格
        if is_table_row(stripped) and i + 1 < len(lines) and is_separator_row(lines[i + 1].strip()):
            header = split_row(stripped)
            rows = []
            j = i + 2
            while j < len(lines) and is_table_row(lines[j].strip()):
                rows.append(split_row(lines[j].strip()))
                j += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Light Grid Accent 1"
            for idx, text in enumerate(header):
                cell = table.rows[0].cells[idx]
                cell.text = ""
                add_inline(cell.paragraphs[0], text, base_bold=True)
            for row in rows:
                cells = table.add_row().cells
                for idx in range(len(header)):
                    cells[idx].text = ""
                    add_inline(cells[idx].paragraphs[0], row[idx] if idx < len(row) else "")
            doc.add_paragraph()
            i = j
            continue

        # 引用块
        if stripped.startswith(">"):
            buf = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                buf.append(lines[i].strip().lstrip(">").strip())
                i += 1
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Pt(18)
            add_inline(para, " ".join(x for x in buf if x), color=RGBColor(0x55, 0x55, 0x55))
            continue

        # 无序列表
        if re.match(r"^[-*]\s+", stripped):
            para = doc.add_paragraph(style="List Bullet")
            add_inline(para, re.sub(r"^[-*]\s+", "", stripped))
            i += 1
            continue

        # 有序列表
        if re.match(r"^\d+[.)]\s+", stripped):
            para = doc.add_paragraph(style="List Number")
            add_inline(para, re.sub(r"^\d+[.)]\s+", "", stripped))
            i += 1
            continue

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_inline(para, stripped)
        i += 1

    doc.save(docx_path)
    print(f"已生成 {docx_path}")


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print(__doc__)
        sys.exit(1)
    build(sys.argv[1], sys.argv[2])
